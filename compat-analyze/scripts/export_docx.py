#!/usr/bin/env python3
"""把 compat-report.md 转成中文 Word 文档。依赖 python-docx（pip install python-docx）。
用法: python export_docx.py compat-report.md compat-report.docx
"""
import sys
from pathlib import Path
from docx import Document
from docx.oxml.ns import qn


def md_to_docx(md_path: str, docx_path: str):
    md = Path(md_path).read_text(encoding="utf-8")
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "SimSun"
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "SimSun")

    for raw in md.splitlines():
        line = raw.rstrip()
        if not line:
            doc.add_paragraph("")
            continue
        if line.startswith("# "):
            p = doc.add_paragraph()
            r = p.add_run(line[2:]); r.bold = True; r.font.size = None
        elif line.startswith("## "):
            p = doc.add_paragraph()
            r = p.add_run(line[3:]); r.bold = True
        elif line.startswith("### "):
            p = doc.add_paragraph()
            r = p.add_run(line[4:]); r.bold = True
        elif line.startswith("- "):
            p = doc.add_paragraph()
            p.style = doc.styles["List Bullet"]
            r = p.add_run(line[2:])
        else:
            p = doc.add_paragraph()
            r = p.add_run(line)
        r.font.name = "SimSun"
        r._element.rPr.rFonts.set(qn("w:eastAsia"), "SimSun")

    doc.save(docx_path)


if __name__ == "__main__":
    if len(sys.argv) != 3:
        print("Usage: export_docx.py <input.md> <output.docx>")
        sys.exit(1)
    md_to_docx(sys.argv[1], sys.argv[2])
